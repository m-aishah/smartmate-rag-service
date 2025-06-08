import httpx
import logging
from typing import List, Tuple
import PyPDF2
from docx import Document
import io
import uuid
from app.config import settings
from app.utils.text_processing import TextChunker

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.text_chunker = TextChunker(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap
        )
        self.max_file_size = settings.max_file_size_mb * 1024 * 1024  # Convert to bytes
    
    async def download_file(self, file_url: str) -> bytes:
        """Download file from URL"""
        try:
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.get(file_url)
                response.raise_for_status()
                
                # Check file size
                content_length = response.headers.get('content-length')
                if content_length and int(content_length) > self.max_file_size:
                    raise ValueError(f"File size exceeds maximum allowed size of {settings.max_file_size_mb}MB")
                
                file_content = response.content
                
                # Double-check actual file size
                if len(file_content) > self.max_file_size:
                    raise ValueError(f"File size exceeds maximum allowed size of {settings.max_file_size_mb}MB")
                
                return file_content
                
        except httpx.RequestError as e:
            logger.error(f"Failed to download file from {file_url}: {e}")
            raise ValueError(f"Failed to download file: {str(e)}")
        except httpx.HTTPStatusError as e:
            logger.error(f"HTTP error downloading file from {file_url}: {e}")
            raise ValueError(f"HTTP error: {e.response.status_code}")
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    continue
            
            full_text = "\n\n".join(text_content)
            
            if not full_text.strip():
                raise ValueError("No text could be extracted from the PDF")
            
            return full_text
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            document = Document(docx_file)
            
            text_content = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n\n".join(text_content)
            
            if not full_text.strip():
                raise ValueError("No text could be extracted from the DOCX")
            
            return full_text
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    async def process_document(self, file_url: str, file_type: str, 
                              document_id: str) -> Tuple[str, List[dict]]:
        """
        Process document and return extracted text and chunks
        Returns: (full_text, chunks_data)
        """
        try:
            # Download file
            logger.info(f"Downloading file from: {file_url}")
            file_content = await self.download_file(file_url)
            
            # Extract text based on file type
            if file_type.lower() == 'pdf':
                text = self.extract_text_from_pdf(file_content)
            elif file_type.lower() == 'docx':
                text = self.extract_text_from_docx(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            logger.info(f"Extracted {len(text)} characters from {file_type} file")
            
            # Chunk the text
            chunks = self.text_chunker.chunk_text(text)
            logger.info(f"Created {len(chunks)} chunks from document")
            
            # Prepare chunks data for database
            chunks_data = []
            for i, chunk in enumerate(chunks):
                chunk_data = {
                    'chunk_id': str(uuid.uuid4()),
                    'document_id': document_id,
                    'content': chunk,
                    'chunk_index': i,
                    'metadata': {
                        'char_count': len(chunk),
                        'word_count': len(chunk.split())
                    }
                }
                chunks_data.append(chunk_data)
            
            return text, chunks_data
            
        except Exception as e:
            logger.error(f"Document processing failed for {document_id}: {e}")
            raise

# Global document processor instance
document_processor = DocumentProcessor()